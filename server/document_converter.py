"""Document conversion tools for MCP Streamable HTTP server."""

import argparse
import base64
import io
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

import uvicorn
from mcp.server.fastmcp import FastMCP

# Third-party libraries for document conversion
try:
    import pypandoc
except ImportError:
    pypandoc = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError:
    canvas = None
    SimpleDocTemplate = None

try:
    import markdown
except ImportError:
    markdown = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
except ImportError:
    DocxDocument = None

# Initialize FastMCP server for Document Conversion tools
# Support both Streamable HTTP and SSE
mcp = FastMCP(name="document-converter", json_response=False, stateless_http=False)

# Supported formats
SUPPORTED_FORMATS = {
    'input': ['txt', 'md', 'html', 'docx', 'pdf', 'rtf'],
    'output': ['txt', 'md', 'html', 'docx', 'pdf', 'rtf']
}


def check_dependencies() -> Dict[str, bool]:
    """Check which conversion libraries are available."""
    return {
        'pypandoc': pypandoc is not None,
        'reportlab': canvas is not None,
        'markdown': markdown is not None,
        'beautifulsoup': BeautifulSoup is not None,
        'python-docx': DocxDocument is not None
    }


def get_file_extension(filename: str) -> str:
    """Get file extension from filename."""
    return Path(filename).suffix.lower().lstrip('.')


def convert_text_to_markdown(text: str) -> str:
    """Convert plain text to markdown format."""
    # Simple text to markdown conversion
    lines = text.split('\n')
    markdown_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            markdown_lines.append('')
        elif line.isupper() and len(line) > 3:
            # Treat all caps as headers
            markdown_lines.append(f'# {line.title()}')
        elif line.endswith(':'):
            # Treat lines ending with colon as subheaders
            markdown_lines.append(f'## {line[:-1]}')
        else:
            markdown_lines.append(line)
    
    return '\n'.join(markdown_lines)


def convert_markdown_to_html(md_text: str) -> str:
    """Convert markdown to HTML."""
    if markdown is None:
        raise ValueError("Markdown library not available. Install with: pip install markdown")
    
    html = markdown.markdown(md_text, extensions=['extra', 'codehilite'])
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #ddd; margin: 0; padding-left: 20px; color: #666; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""


def convert_html_to_text(html_content: str) -> str:
    """Convert HTML to plain text."""
    if BeautifulSoup is None:
        raise ValueError("BeautifulSoup library not available. Install with: pip install beautifulsoup4")
    
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text(separator='\n', strip=True)


def convert_to_pdf(content: str, content_type: str = 'text') -> bytes:
    """Convert content to PDF format."""
    if SimpleDocTemplate is None:
        raise ValueError("ReportLab library not available. Install with: pip install reportlab")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    if content_type == 'html':
        # Simple HTML to text conversion for PDF
        if BeautifulSoup:
            soup = BeautifulSoup(content, 'html.parser')
            content = soup.get_text(separator='\n', strip=True)
    
    # Split content into paragraphs
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            # Check if it looks like a header (short line, all caps, etc.)
            if len(para.strip()) < 50 and (para.isupper() or para.startswith('#')):
                p = Paragraph(para.strip().lstrip('#').strip(), styles['Heading1'])
            else:
                p = Paragraph(para.strip(), styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def convert_to_docx(content: str, content_type: str = 'text') -> bytes:
    """Convert content to DOCX format."""
    if DocxDocument is None:
        raise ValueError("python-docx library not available. Install with: pip install python-docx")
    
    doc = DocxDocument()
    
    if content_type == 'html':
        # Simple HTML to text conversion for DOCX
        if BeautifulSoup:
            soup = BeautifulSoup(content, 'html.parser')
            content = soup.get_text(separator='\n', strip=True)
    
    # Split content into paragraphs
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            # Check if it looks like a header
            if len(para.strip()) < 50 and (para.isupper() or para.startswith('#')):
                heading = doc.add_heading(para.strip().lstrip('#').strip(), level=1)
            else:
                doc.add_paragraph(para.strip())
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@mcp.tool()
async def convert_document(content: str, input_format: str, output_format: str, filename: Optional[str] = None) -> str:
    """Convert document from one format to another.
    
    Args:
        content: The document content (text or base64 encoded for binary formats)
        input_format: Input format (txt, md, html, docx, pdf, rtf)
        output_format: Output format (txt, md, html, docx, pdf, rtf)
        filename: Optional filename for context
    
    Returns:
        Converted document content (text or base64 encoded for binary formats)
    """
    input_format = input_format.lower().strip()
    output_format = output_format.lower().strip()
    
    # Validate formats
    if input_format not in SUPPORTED_FORMATS['input']:
        return f"Error: Unsupported input format '{input_format}'. Supported: {', '.join(SUPPORTED_FORMATS['input'])}"
    
    if output_format not in SUPPORTED_FORMATS['output']:
        return f"Error: Unsupported output format '{output_format}'. Supported: {', '.join(SUPPORTED_FORMATS['output'])}"
    
    try:
        # Handle input content
        if input_format in ['pdf', 'docx']:
            # For binary formats, expect base64 encoded content
            try:
                binary_content = base64.b64decode(content)
            except Exception:
                return "Error: Binary formats require base64 encoded content"
        else:
            # Text formats
            text_content = content
        
        # Convert based on input format
        if input_format == 'txt':
            working_content = content
            working_type = 'text'
        elif input_format == 'md':
            working_content = content
            working_type = 'markdown'
        elif input_format == 'html':
            working_content = content
            working_type = 'html'
        elif input_format == 'docx':
            # Extract text from DOCX (simplified)
            if DocxDocument:
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    tmp.write(binary_content)
                    tmp.flush()
                    doc = DocxDocument(tmp.name)
                    working_content = '\n'.join([para.text for para in doc.paragraphs])
                    working_type = 'text'
                os.unlink(tmp.name)
            else:
                return "Error: python-docx library not available for DOCX processing"
        elif input_format == 'pdf':
            # PDF text extraction would require additional libraries like PyPDF2
            return "Error: PDF input processing not yet implemented. Please use pypandoc or similar tools."
        else:
            working_content = content
            working_type = 'text'
        
        # Convert to output format
        if output_format == 'txt':
            if working_type == 'html':
                result = convert_html_to_text(working_content)
            elif working_type == 'markdown':
                # Simple markdown to text conversion
                result = working_content.replace('#', '').replace('*', '').replace('_', '')
            else:
                result = working_content
            return result
        
        elif output_format == 'md':
            if working_type == 'text':
                result = convert_text_to_markdown(working_content)
            elif working_type == 'html':
                # HTML to markdown conversion (simplified)
                text = convert_html_to_text(working_content)
                result = convert_text_to_markdown(text)
            else:
                result = working_content
            return result
        
        elif output_format == 'html':
            if working_type == 'markdown':
                result = convert_markdown_to_html(working_content)
            elif working_type == 'text':
                # Text to HTML
                escaped_content = working_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                paragraphs = escaped_content.split('\n\n')
                html_paragraphs = [f'<p>{para.replace(chr(10), "<br>")}</p>' for para in paragraphs if para.strip()]
                result = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Document</title>
    <style>body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}</style>
</head>
<body>
{''.join(html_paragraphs)}
</body>
</html>"""
            else:
                result = working_content
            return result
        
        elif output_format == 'pdf':
            pdf_bytes = convert_to_pdf(working_content, working_type)
            return base64.b64encode(pdf_bytes).decode('utf-8')
        
        elif output_format == 'docx':
            docx_bytes = convert_to_docx(working_content, working_type)
            return base64.b64encode(docx_bytes).decode('utf-8')
        
        else:
            return f"Error: Output format '{output_format}' conversion not implemented"
    
    except Exception as e:
        return f"Error during conversion: {str(e)}"


@mcp.tool()
async def list_supported_formats() -> str:
    """List all supported input and output formats for document conversion.
    
    Returns:
        Information about supported formats and available libraries
    """
    deps = check_dependencies()
    
    result = "Document Converter - Supported Formats\n"
    result += "=" * 40 + "\n\n"
    
    result += "Input Formats: " + ", ".join(SUPPORTED_FORMATS['input']) + "\n"
    result += "Output Formats: " + ", ".join(SUPPORTED_FORMATS['output']) + "\n\n"
    
    result += "Library Dependencies:\n"
    result += "-" * 20 + "\n"
    for lib, available in deps.items():
        status = "✓ Available" if available else "✗ Not installed"
        result += f"{lib}: {status}\n"
    
    result += "\nInstallation Commands:\n"
    result += "-" * 20 + "\n"
    result += "pip install pypandoc markdown beautifulsoup4 reportlab python-docx\n\n"
    
    result += "Notes:\n"
    result += "- Binary formats (PDF, DOCX) require base64 encoding\n"
    result += "- Some conversions may require additional system dependencies\n"
    result += "- For advanced PDF processing, consider installing PyPDF2 or pdfplumber\n"
    
    return result


@mcp.tool()
async def convert_file_batch(files: List[Dict[str, str]], output_format: str) -> str:
    """Convert multiple files to the same output format.
    
    Args:
        files: List of dictionaries with 'content', 'input_format', and optional 'filename'
        output_format: Target output format for all files
    
    Returns:
        Results of batch conversion
    """
    results = []
    
    for i, file_info in enumerate(files):
        content = file_info.get('content', '')
        input_format = file_info.get('input_format', 'txt')
        filename = file_info.get('filename', f'file_{i+1}')
        
        try:
            converted = await convert_document(content, input_format, output_format, filename)
            results.append({
                'filename': filename,
                'status': 'success',
                'content': converted
            })
        except Exception as e:
            results.append({
                'filename': filename,
                'status': 'error',
                'error': str(e)
            })
    
    # Format results
    output = f"Batch Conversion Results ({len(files)} files to {output_format})\n"
    output += "=" * 50 + "\n\n"
    
    for result in results:
        output += f"File: {result['filename']}\n"
        output += f"Status: {result['status']}\n"
        if result['status'] == 'error':
            output += f"Error: {result['error']}\n"
        else:
            content_preview = result['content'][:100] + "..." if len(result['content']) > 100 else result['content']
            output += f"Content Preview: {content_preview}\n"
        output += "-" * 30 + "\n"
    
    return output


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Run MCP Document Converter Server")
    parser.add_argument("--port", type=int, default=8124, help="Port to listen on")
    parser.add_argument("--host", type=str, default="localhost", help="Host to bind to")
    args = parser.parse_args()
    
    print(f"Starting Document Converter MCP Server on {args.host}:{args.port}")
    print("Supported transports: Streamable HTTP, SSE")
    
    # Check dependencies on startup
    deps = check_dependencies()
    print("\nLibrary Status:")
    for lib, available in deps.items():
        status = "✓" if available else "✗"
        print(f"  {status} {lib}")
    
    # Start the server with Streamable HTTP transport
    uvicorn.run(mcp.streamable_http_app, host=args.host, port=args.port)